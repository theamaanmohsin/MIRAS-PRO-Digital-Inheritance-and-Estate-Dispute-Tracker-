#!/usr/bin/env python3
"""
Build MirasPro PROJECT_MANUAL.md into a styled, interactive Word document.
Requires: python-docx, requests (optional for diagram images)
"""
from __future__ import annotations

import base64
import re
import sys
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.shared import Emu

# MirasPro brand (miras-theme.css)
C_PRIMARY = RGBColor(0, 106, 255)
C_PRIMARY_DARK = RGBColor(0, 86, 214)
C_SECONDARY = RGBColor(31, 41, 55)
C_TEXT = RGBColor(17, 24, 39)
C_MUTED = RGBColor(100, 116, 139)
C_SUCCESS = RGBColor(21, 128, 61)
C_WARNING = RGBColor(217, 119, 6)
C_SURFACE = RGBColor(248, 250, 252)
C_WHITE = RGBColor(255, 255, 255)
C_TABLE_HEADER_BG = "006AFF"
C_TABLE_ALT_BG = "F8FAFC"
C_CALLOUT_BG = "EAF2FF"

ROOT = Path(__file__).resolve().parent.parent
MD_PATH = ROOT / "PROJECT_MANUAL.md"
OUT_PATH = ROOT / "MirasPro_Project_Manual.docx"
DIAGRAM_DIR = ROOT / "Scripts" / "diagram_cache"


def rgb_hex(rgb: RGBColor) -> str:
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


def set_cell_shading(cell, fill_hex: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def set_paragraph_shading(paragraph, fill_hex: str):
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    p_pr.append(shading)


def add_bookmark(paragraph, bookmark_id: int, name: str):
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_internal_link(paragraph, text: str, bookmark_name: str, color: RGBColor = C_PRIMARY):
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bookmark_name)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), rgb_hex(color))
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(color_el)
    r_pr.append(u)
    new_run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def insert_word_toc(paragraph):
    """Insert Word TOC field — right-click Update Field in Word to refresh."""
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click here → Update Field to refresh Table of Contents"
    run2 = OxmlElement("w:r")
    run2.append(placeholder)
    paragraph._p.append(run2)
    run._r.append(fld_end)


def fetch_mermaid_png(mermaid_code: str, name: str) -> Path | None:
    try:
        import urllib.request

        encoded = base64.urlsafe_b64encode(mermaid_code.encode("utf-8")).decode().rstrip("=")
        url = f"https://mermaid.ink/img/{encoded}?type=png&bgColor=!f8fafc"
        DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
        out = DIAGRAM_DIR / f"{name}.png"
        req = urllib.request.Request(url, headers={"User-Agent": "MirasPro-Manual-Builder/1.0"})
        with urllib.request.urlopen(req, timeout=60) as resp:
            data = resp.read()
        if len(data) < 500:
            return None
        out.write_bytes(data)
        return out
    except Exception as e:
        print(f"  [diagram] {name}: skip ({e})")
        return None


def make_anchor(title: str, used: set[str]) -> str:
    a = re.sub(r"[^\w\s-]", "", title.lower())
    a = re.sub(r"\s+", "_", a.strip())[:50] or "section"
    if a not in used:
        used.add(a)
        return a
    i = 2
    while f"{a}_{i}" in used:
        i += 1
    a = f"{a}_{i}"
    used.add(a)
    return a


def extract_headings(md: str) -> list[tuple[int, str, str]]:
    used: set[str] = set()
    entries: list[tuple[int, str, str]] = []
    for line in md.splitlines():
        if line.startswith("### "):
            t = line[4:].strip()
            entries.append((3, t, make_anchor(t, used)))
        elif line.startswith("## ") and "Table of Contents" not in line:
            t = line[3:].strip()
            level = 1 if re.match(r"^\d+\.", t) else 2
            entries.append((level, t, make_anchor(t, used)))
    return entries


def parse_inline_runs(paragraph, text: str, base_bold=False, base_italic=False):
    """Parse **bold** and `code` inline markers."""
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
            if base_italic:
                r.italic = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9)
            r.font.color.rgb = C_PRIMARY_DARK
        else:
            r = paragraph.add_run(part)
            r.bold = base_bold
            r.italic = base_italic


class ManualBuilder:
    def __init__(self, heading_plan: list[tuple[int, str, str]]):
        self.doc = Document()
        self.bookmark_id = 0
        self.toc_entries = heading_plan
        self._heading_queue = list(heading_plan)
        self.diagram_idx = 0
        self._setup_styles()
        self._setup_page()

    def _setup_page(self):
        for section in self.doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.2)
            section.right_margin = Cm(2.2)
            header = section.header
            hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            hp.clear()
            r1 = hp.add_run("MirasPro")
            r1.bold = True
            r1.font.color.rgb = C_PRIMARY
            r1.font.size = Pt(11)
            r1.font.name = "Segoe UI"
            r2 = hp.add_run("  |  Inheritance System — Project Manual")
            r2.font.size = Pt(9)
            r2.font.color.rgb = C_MUTED
            r2.font.name = "Segoe UI"
            hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            footer = section.footer
            fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            fp.clear()
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._add_page_number_field(fp)

    def _add_page_number_field(self, paragraph):
        run = paragraph.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_sep)
        t = OxmlElement("w:t")
        t.text = "1"
        run2 = OxmlElement("w:r")
        run2.append(t)
        paragraph._p.append(run2)
        run._r.append(fld_end)
        r = paragraph.add_run("   ·   Confidential — Internal Use")
        r.font.size = Pt(8)
        r.font.color.rgb = C_MUTED

    def _setup_styles(self):
        normal = self.doc.styles["Normal"]
        normal.font.name = "Segoe UI"
        normal.font.size = Pt(11)
        normal.font.color.rgb = C_TEXT
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        normal.paragraph_format.line_spacing = 1.15

        for i, size, color in [
            (1, 22, C_SECONDARY),
            (2, 16, C_PRIMARY),
            (3, 13, C_PRIMARY_DARK),
        ]:
            style = self.doc.styles[f"Heading {i}"]
            style.font.name = "Segoe UI Semibold"
            style.font.size = Pt(size)
            style.font.color.rgb = color
            style.font.bold = True
            style.paragraph_format.space_before = Pt(14 if i > 1 else 0)
            style.paragraph_format.space_after = Pt(8)
            style.paragraph_format.keep_with_next = True

    def add_cover(self):
        # Cover block
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(72)
        set_paragraph_shading(p, C_TABLE_HEADER_BG)
        r = p.add_run("م")
        r.font.size = Pt(36)
        r.font.color.rgb = C_WHITE

        t = self.doc.add_paragraph()
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_shading(t, C_TABLE_HEADER_BG)
        r = t.add_run("MirasPro")
        r.bold = True
        r.font.size = Pt(42)
        r.font.color.rgb = C_WHITE
        r.font.name = "Segoe UI"

        sub = self.doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_shading(sub, C_TABLE_HEADER_BG)
        r = sub.add_run("Inheritance System")
        r.font.size = Pt(20)
        r.font.color.rgb = RGBColor(219, 234, 254)

        doc_title = self.doc.add_paragraph()
        doc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_shading(doc_title, rgb_hex(C_SECONDARY))
        doc_title.paragraph_format.space_before = Pt(24)
        doc_title.paragraph_format.space_after = Pt(24)
        r = doc_title.add_run("PROJECT MANUAL")
        r.bold = True
        r.font.size = Pt(18)
        r.font.color.rgb = C_WHITE
        r.font.letter_spacing = Pt(2)

        meta = self.doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.paragraph_format.space_before = Pt(36)
        for line in [
            "Sharia-Compliant Digital Inheritance & Property Management",
            "ASP.NET Core Blazor Server · Entity Framework Core · SQL Server",
            "Version 1.0  ·  May 2026",
        ]:
            r = meta.add_run(line + "\n")
            r.font.size = Pt(11)
            r.font.color.rgb = C_MUTED

        # Quick navigation callout
        self.doc.add_page_break()
        box = self.doc.add_paragraph()
        box.paragraph_format.space_before = Pt(12)
        set_paragraph_shading(box, C_CALLOUT_BG)
        r = box.add_run("📌  How to use this document\n")
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = C_PRIMARY
        tips = box.add_run(
            "• Click any blue link in the Navigation Index to jump to a section.\n"
            "• Use Word’s Table of Contents (next page) — right-click → Update Field.\n"
            "• Diagrams are embedded as figures; zoom or print for clarity.\n"
            "• Press Ctrl+F to search routes, tables, or Faraid terms."
        )
        tips.font.size = Pt(10)
        tips.font.color.rgb = C_TEXT

    def add_navigation_index(self):
        self.doc.add_heading("Navigation Index", level=1)
        intro = self.doc.add_paragraph(
            "Interactive links below jump to major sections. Subsections appear in the auto-generated Table of Contents on the following page."
        )
        intro.runs[0].font.color.rgb = C_MUTED
        intro.runs[0].font.size = Pt(10)

        for level, title, anchor in self.toc_entries:
            if level > 2:
                continue
            p = self.doc.add_paragraph()
            indent = "    " * (level - 1)
            prefix = p.add_run(indent)
            prefix.font.size = Pt(10)
            num = re.match(r"^(\d+(?:\.\d+)?)", title)
            display = title if num else title
            add_internal_link(p, display, anchor)
            if level == 1:
                p.paragraph_format.space_before = Pt(6)

        self.doc.add_page_break()
        self.doc.add_heading("Table of Contents (Auto-Update)", level=1)
        note = self.doc.add_paragraph(
            "In Microsoft Word: right-click the box below → Update Field → Update entire table."
        )
        note.runs[0].italic = True
        note.runs[0].font.size = Pt(9)
        note.runs[0].font.color.rgb = C_MUTED
        toc_p = self.doc.add_paragraph()
        insert_word_toc(toc_p)
        self.doc.add_page_break()

    def add_heading_with_bookmark(self, text: str, level: int):
        clean = re.sub(r"^#+\s*", "", text).strip()
        anchor = clean
        if self._heading_queue and self._heading_queue[0][1] == clean:
            _, _, anchor = self._heading_queue.pop(0)
        elif self._heading_queue:
            _, _, anchor = self._heading_queue.pop(0)

        h = self.doc.add_heading(clean, level=level)
        self.bookmark_id += 1
        add_bookmark(h, self.bookmark_id, anchor)

        if level == 1:
            self.doc.add_paragraph().paragraph_format.space_after = Pt(0)

    def add_callout(self, text: str, kind: str = "info"):
        p = self.doc.add_paragraph()
        colors = {
            "info": (C_CALLOUT_BG, C_PRIMARY, "ℹ️ "),
            "tip": ("E8F5E9", C_SUCCESS, "✓ "),
            "warn": ("FFF7ED", C_WARNING, "⚠ "),
        }
        bg, fg, icon = colors.get(kind, colors["info"])
        set_paragraph_shading(p, bg)
        r = p.add_run(icon)
        r.bold = True
        parse_inline_runs(p, text)
        for run in p.runs[1:]:
            run.font.size = Pt(10)

    def add_code_block(self, lines: list[str], lang: str = ""):
        if lang == "mermaid":
            self.diagram_idx += 1
            code = "\n".join(lines)
            name = f"diagram_{self.diagram_idx}"
            png = fetch_mermaid_png(code, name)
            cap = self.doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cap.add_run(f"Figure {self.diagram_idx}: UML Diagram")
            r.bold = True
            r.font.color.rgb = C_PRIMARY
            r.font.size = Pt(10)
            if png and png.exists():
                pic = self.doc.add_paragraph()
                pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = pic.add_run()
                run.add_picture(str(png), width=Inches(6.2))
                self.add_callout(
                    "Interactive diagram — refer to this figure when following related workflows in the text.",
                    "tip",
                )
            else:
                self.add_callout(
                    "Diagram could not be loaded offline. See PROJECT_MANUAL.md for Mermaid source or connect to the internet and rebuild.",
                    "warn",
                )
                p = self.doc.add_paragraph()
                set_paragraph_shading(p, "F1F5F9")
                for line in lines[:12]:
                    r = p.add_run(line + "\n")
                    r.font.name = "Consolas"
                    r.font.size = Pt(8)
                if len(lines) > 12:
                    p.add_run(f"... ({len(lines) - 12} more lines)\n").font.size = Pt(8)
            return

        p = self.doc.add_paragraph()
        set_paragraph_shading(p, "F1F5F9")
        p.paragraph_format.left_indent = Inches(0.15)
        for line in lines:
            r = p.add_run(line + "\n")
            r.font.name = "Consolas"
            r.font.size = Pt(9)
            r.font.color.rgb = C_SECONDARY

    def add_table_from_rows(self, rows: list[list[str]]):
        if not rows:
            return
        cols = max(len(r) for r in rows)
        table = self.doc.add_table(rows=len(rows), cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        for i, row in enumerate(rows):
            for j in range(cols):
                cell = table.rows[i].cells[j]
                cell.text = ""
                text = row[j] if j < len(row) else ""
                p = cell.paragraphs[0]
                parse_inline_runs(p, text.strip())
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Segoe UI"
                if i == 0:
                    set_cell_shading(cell, C_TABLE_HEADER_BG)
                    for run in p.runs:
                        run.bold = True
                        run.font.color.rgb = C_WHITE
                elif i % 2 == 0:
                    set_cell_shading(cell, C_TABLE_ALT_BG)
        self.doc.add_paragraph()

    def add_bullet(self, text: str, ordered: bool = False, number: int | None = None):
        style = "List Number" if ordered else "List Bullet"
        p = self.doc.add_paragraph(style=style)
        if ordered and number:
            p.text = ""
            r = p.add_run(f"{number}. ")
            r.bold = True
            r.font.color.rgb = C_PRIMARY
        parse_inline_runs(p, text.lstrip("- ").strip())

    def build_from_markdown(self, md: str):
        lines = md.splitlines()
        i = 0
        in_code = False
        code_lang = ""
        code_lines: list[str] = []
        table_rows: list[list[str]] = []
        in_table = False
        skip_until = 0  # skip duplicate TOC in md

        while i < len(lines):
            line = lines[i]

            # Skip markdown TOC block (we build our own)
            if line.strip() == "## Table of Contents":
                i += 1
                while i < len(lines) and not lines[i].startswith("## 1."):
                    i += 1
                continue

            if line.strip() == "---":
                i += 1
                continue

            if line.startswith("```"):
                if not in_code:
                    in_code = True
                    code_lang = line[3:].strip()
                    code_lines = []
                else:
                    in_code = False
                    self.add_code_block(code_lines, code_lang)
                    code_lines = []
                    code_lang = ""
                i += 1
                continue

            if in_code:
                code_lines.append(line)
                i += 1
                continue

            if line.startswith("|") and "|" in line[1:]:
                if re.match(r"^\|[\s\-:|]+\|$", line):
                    i += 1
                    continue
                cells = [c.strip() for c in line.strip("|").split("|")]
                table_rows.append(cells)
                in_table = True
                i += 1
                continue
            elif in_table:
                self.add_table_from_rows(table_rows)
                table_rows = []
                in_table = False

            if line.startswith("# ") and not line.startswith("## "):
                i += 1
                continue  # skip md title (cover handles it)

            if line.startswith("### "):
                self.add_heading_with_bookmark(line[4:].strip(), 3)
                i += 1
                continue
            if line.startswith("## "):
                t = line[3:].strip()
                lvl = 1 if re.match(r"^\d+\.", t) else 2
                self.add_heading_with_bookmark(t, lvl)
                i += 1
                continue

            if line.startswith("- ") or line.startswith("* "):
                self.add_bullet(line)
                i += 1
                continue

            m = re.match(r"^(\d+)\.\s+(.*)", line)
            if m:
                self.add_bullet(m.group(2), ordered=True, number=int(m.group(1)))
                i += 1
                continue

            if not line.strip():
                i += 1
                continue

            p = self.doc.add_paragraph()
            parse_inline_runs(p, line.strip())

            i += 1

        if in_table and table_rows:
            self.add_table_from_rows(table_rows)

    def build(self):
        md = MD_PATH.read_text(encoding="utf-8")
        self._heading_queue = list(self.toc_entries)
        self.add_cover()
        self.add_navigation_index()
        self.build_from_markdown(md)

        # Appendix callout
        self.doc.add_page_break()
        end = self.doc.add_paragraph()
        end.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_shading(end, C_TABLE_HEADER_BG)
        r = end.add_run("— End of MirasPro Project Manual —")
        r.font.color.rgb = C_WHITE
        r.bold = True

        self.doc.save(OUT_PATH)
        print(f"Created: {OUT_PATH}")


def main():
    if not MD_PATH.exists():
        print(f"Missing {MD_PATH}")
        sys.exit(1)
    md = MD_PATH.read_text(encoding="utf-8")
    plan = extract_headings(md)
    builder = ManualBuilder(plan)
    builder.build()


if __name__ == "__main__":
    main()
